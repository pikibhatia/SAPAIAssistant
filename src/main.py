import streamlit as st
from streamlit_option_menu import option_menu
import os
import requests
import re
import io
from typing import Dict, List, Optional, Union, Any, BinaryIO
import base64
from docx import Document
import tempfile
import zipfile
import fitz  # PyMuPDF
import autogen
import openai

###################
# Document Processor
###################

class DocumentProcessor:
    """
    Utility for extracting text from various document formats.
    Supports .txt, .docx, and .pdf files.
    """
    
    @staticmethod
    def extract_text(file: Union[BinaryIO, io.BytesIO]) -> str:
        """
        Extract text from the uploaded file based on its extension
        
        Args:
            file: File object from streamlit uploader
            
        Returns:
            Extracted text content
        """
        # Get file extension from name
        file_name = file.name.lower() if hasattr(file, 'name') else 'unknown.txt'
        
        if file_name.endswith('.txt'):
            return DocumentProcessor._process_txt(file)
        elif file_name.endswith('.docx'):
            return DocumentProcessor._process_docx(file)
        elif file_name.endswith('.pdf'):
            return DocumentProcessor._process_pdf(file)
        else:
            # For any other format, try to read as text
            return DocumentProcessor._process_txt(file)
    
    @staticmethod
    def _process_txt(file: Union[BinaryIO, io.BytesIO]) -> str:
        """Process text files."""
        try:
            # Reset file pointer to the beginning
            file.seek(0)
            
            # Read and decode text content
            text_content = file.read().decode('utf-8')
            return text_content
        except UnicodeDecodeError:
            # If UTF-8 decode fails, try with another encoding
            file.seek(0)
            text_content = file.read().decode('latin-1')
            return text_content
    
    @staticmethod
    def _process_docx(file: Union[BinaryIO, io.BytesIO]) -> str:
        """Process Word documents."""
        try:
            # Load the document from binary content
            doc = Document(file)
            
            # Extract all paragraphs
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text)
                    full_text.append(" | ".join(row_text))
            
            return "\n".join(full_text)
        except Exception as e:
            raise Exception(f"Error processing DOCX file: {str(e)}")
    
    @staticmethod
    def _process_pdf(file: Union[BinaryIO, io.BytesIO]) -> str:
        """Process PDF documents."""
        try:
            # Load PDF from binary content
            file_bytes = file.read()
            
            # Create a temporary file to work with PyMuPDF
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                temp_file.write(file_bytes)
                temp_path = temp_file.name
            
            # Extract text using PyMuPDF
            doc = fitz.open(temp_path)
            text_content = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text_content.append(page.get_text())
            
            # Clean up the temporary file
            os.unlink(temp_path)
            
            return "\n".join(text_content)
        except Exception as e:
            raise Exception(f"Error processing PDF file: {str(e)}")

###################
# File Handler
###################

class FileHandler:
    """
    Utility for handling file operations including:
    - Creating downloadable files
    - Generating .docx documents
    - Creating .zip archives
    """
    
    @staticmethod
    def get_download_link(content: str, filename: str, label: str) -> str:
        """
        Generate a download link for text content
        
        Args:
            content: Text content to download
            filename: Name of the file
            label: Link text to display
            
        Returns:
            HTML string with download link
        """
        # Convert string to bytes and encode as base64
        b64 = base64.b64encode(content.encode()).decode()
        
        # Create the HTML download link
        href = f'<a href="data:file/txt;base64,{b64}" download="{filename}">{label}</a>'
        return href
    
    @staticmethod
    def create_docx_download(content: str, filename: str = "document.docx") -> bytes:
        """
        Create a Word document from text content
        
        Args:
            content: Text content to include in the document
            filename: Name of the output file
            
        Returns:
            Word document as bytes
        """
        # Create a new Document
        doc = Document()
        
        # Add a title (first line)
        first_line = content.split('\n', 1)[0].strip().replace('#', '').strip()
        if first_line:
            doc.add_heading(first_line, level=1)
        
        # Process the content
        # Split by markdown headings and add appropriate Word headings
        content_blocks = re.split(r'(#+\s+.*)', content)
        
        for block in content_blocks:
            block = block.strip()
            if not block:
                continue
                
            # Check if the block is a heading
            if block.startswith('#'):
                heading_match = re.match(r'(#+)\s+(.*)', block)
                if heading_match:
                    level = min(len(heading_match.group(1)), 9)  # Word supports up to 9 levels
                    text = heading_match.group(2).strip()
                    doc.add_heading(text, level=level)
            else:
                # Add paragraphs
                paragraphs = block.split('\n')
                for p in paragraphs:
                    if p.strip():
                        doc.add_paragraph(p.strip())
        
        # Save the document to a bytes buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
    
    @staticmethod
    def create_zip_download(files: Dict[str, str], zip_filename: str = "output.zip") -> bytes:
        """
        Create a zip archive containing multiple files
        
        Args:
            files: Dictionary with filenames as keys and file contents as values
            zip_filename: Name of the output zip file
            
        Returns:
            Zip archive as bytes
        """
        # Create a BytesIO object to store the zip file
        zip_buffer = io.BytesIO()
        
        # Create the zip file
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for filename, content in files.items():
                # Handle nested directory paths in the filename
                if '/' in filename:
                    # Extract the directory part
                    directory = filename.rsplit('/', 1)[0]
                    
                    # Add directory entries if needed
                    parts = directory.split('/')
                    for i in range(len(parts)):
                        dir_path = '/'.join(parts[:i+1])
                        # Try to create directory (will fail silently if it exists)
                        try:
                            zip_info = zipfile.ZipInfo(dir_path + '/')
                            zip_file.writestr(zip_info, '')
                        except:
                            pass
                
                # Add the file to the zip
                if isinstance(content, str):
                    # For text content
                    zip_file.writestr(filename, content)
                else:
                    # For binary content
                    zip_file.writestr(filename, content)
        
        # Seek to the beginning of the buffer
        zip_buffer.seek(0)
        
        return zip_buffer.getvalue()
    
    @staticmethod
    def display_download_button(content: Union[str, bytes], filename: str, label: str, 
                               mime: str = "text/plain") -> None:
        """
        Display a download button in Streamlit
        
        Args:
            content: File content (string or bytes)
            filename: Name of the file to download
            label: Text to display on the button
            mime: MIME type of the file
        """
        if isinstance(content, str):
            content = content.encode()
            
        st.download_button(
            label=label,
            data=content,
            file_name=filename,
            mime=mime
        )

###################
# GROQ Client
###################

class GroqClient:
    """Client for interacting with GROQ API."""
    
    def __init__(self):
        # First try to get from session state, then fallback to environment
        if hasattr(st, 'session_state') and 'groq_api_key' in st.session_state and st.session_state.groq_api_key:
            self.api_key = st.session_state.groq_api_key
        else:
            self.api_key = os.getenv("GROQ_API_KEY", "")
            
        if not self.api_key:
            raise ValueError("GROQ API key not found. Please provide your API key in the sidebar.")
        
        self.base_url = "https://api.groq.com/openai/v1"
        self.headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        # Default to best available model
        self.model = "llama-3.3-70b-versatile"
    
    def set_model(self, model_name: str) -> None:
        """Set the model to use for completions."""
        self.model = model_name
    
    def chat_completion(
        self, 
        messages: List[Dict[str, str]], 
        temperature: float = 0.7,
        max_tokens: int = 4000,
        stream: bool = False
    ) -> Union[Dict[str, Any], Any]:
        """
        Send a chat completion request to GROQ API.
        
        Args:
            messages: List of message objects with role and content
            temperature: Sampling temperature (0.0-1.0)
            max_tokens: Maximum tokens to generate
            stream: Whether to stream the response
            
        Returns:
            Response from GROQ API
        """
        url = f"{self.base_url}/chat/completions"
        
        payload = {
            "model": self.model,
            "messages": messages,
            "temperature": temperature,
            "max_tokens": max_tokens,
            "stream": stream
        }
        
        try:
            response = requests.post(url, headers=self.headers, json=payload)
            response.raise_for_status()
            
            if stream:
                # Return the response object for streaming
                return response
            else:
                return response.json()
                
        except requests.exceptions.RequestException as e:
            error_msg = f"Error calling GROQ API: {str(e)}"
            if hasattr(e, 'response') and e.response is not None:
                try:
                    error_data = e.response.json()
                    error_msg += f" - {error_data.get('error', {}).get('message', '')}"
                except:
                    error_msg += f" - Status code: {e.response.status_code}"
            
            raise Exception(error_msg)

    def extract_content(self, response: Dict[str, Any]) -> str:
        """Extract content from GROQ API response."""
        try:
            return response["choices"][0]["message"]["content"]
        except (KeyError, IndexError) as e:
            raise Exception(f"Error extracting content from GROQ response: {str(e)}")

    def get_available_models(self) -> List[str]:
        """Get list of available models from GROQ API."""
        url = f"{self.base_url}/models"
        
        try:
            response = requests.get(url, headers=self.headers)
            response.raise_for_status()
            models_data = response.json()
            
            # Extract model IDs
            model_ids = [model["id"] for model in models_data.get("data", [])]
            return model_ids
            
        except requests.exceptions.RequestException as e:
            # Return default models if API call fails
            return ["llama-3.3-70b-versatile", "llama-3.3-70b-versatile", "llama-3.3-70b-versatile"]

###################
# Agent Class
###################

class Agent:
    """Base class for all agents."""
    
    def __init__(self, groq_client: GroqClient):
        self.groq_client = groq_client
        self.thinking_steps = []
    
    def _add_thinking_step(self, step: str) -> None:
        """Add a thinking step to the log."""
        self.thinking_steps.append(step)
    
    def get_thinking_steps(self) -> List[str]:
        """Get the agent's thinking steps for explanation."""
        return self.thinking_steps
    
    def _create_chat_completion(self, messages, temperature=0.7) -> str:
        """Create a chat completion with the GROQ API."""
        response = self.groq_client.chat_completion(
            messages=messages,
            temperature=temperature
        )
        return self.groq_client.extract_content(response)

###################
# FS Analyzer Agent
###################

class FSAnalyzerAgent(Agent):
    """
    Agent that analyzes SAP Functional Specifications and extracts key information.
    """
    
    def analyze_fs(self, fs_text: str) -> str:
        """
        Analyze a functional specification and extract key points.
        
        Args:
            fs_text: The functional specification text
            
        Returns:
            Analysis results
        """
        self._add_thinking_step("Creating a structured task plan for analyzing the functional specification")
        task_plan = self._create_task_plan(fs_text)
        
        self._add_thinking_step("Identifying any potential information gaps in the specification")
        info_gaps = self._identify_information_gaps(fs_text)
        
        self._add_thinking_step("Performing detailed analysis of the functional specification")
        analysis = self._perform_analysis(fs_text)
        
        # Combine the results
        full_analysis = f"""
# SAP Functional Specification Analysis

## Task Planning
{task_plan}

## Information Gaps
{info_gaps}

## Detailed Analysis
{analysis}
"""
        
        return full_analysis
    
    def _create_task_plan(self, fs_text: str) -> str:
        """Create a structured task plan for processing the FS."""
        prompt = """
As an SAP functional specification analyzer, create a structured implementation plan based on the following functional specification. 
Format your response as a markdown list with the following sections:
1. Key SAP Objects Required
2. Development Tasks in Priority Order
3. Configuration Tasks
4. Testing Requirements

Here is the Functional Specification:
"""
        
        messages = [
            {"role": "system", "content": "You are an SAP expert specializing in analyzing functional specifications."},
            {"role": "user", "content": f"{prompt}\n\n{fs_text}"}
        ]
        
        return self._create_chat_completion(messages)
    
    def _identify_information_gaps(self, fs_text: str) -> str:
        """Identify missing information in the functional specification."""
        prompt = """
Review the following SAP functional specification and identify any missing information or ambiguities that should be clarified before development.
Focus on:
- Missing technical details
- Unclear requirements
- Integration points without sufficient details
- Security considerations
- Performance expectations
- Edge cases not addressed

Format your response as a markdown list of questions or clarifications needed.

Functional Specification:
"""
        
        messages = [
            {"role": "system", "content": "You are an SAP expert specializing in quality assessment of functional specifications."},
            {"role": "user", "content": f"{prompt}\n\n{fs_text}"}
        ]
        
        return self._create_chat_completion(messages)
    
    def _perform_analysis(self, fs_text: str) -> str:
        """Perform detailed analysis of the functional specification."""
        prompt = """
Perform a detailed analysis of the following SAP functional specification.
Include:
1. Business process mapping and impact
2. Technical component identification
3. Data flow analysis
4. Integration requirements
5. Performance considerations
6. Security implications

Format your analysis in markdown with clear sections.

Functional Specification:
"""
        
        messages = [
            {"role": "system", "content": "You are an SAP expert specializing in detailed analysis of functional specifications."},
            {"role": "user", "content": f"{prompt}\n\n{fs_text}"}
        ]
        
        return self._create_chat_completion(messages)

###################
# TS Generator Agent
###################

class TSGeneratorAgent(Agent):
    """
    Agent that generates SAP Technical Specifications from Functional Specifications.
    """
    
    def generate_ts(self, fs_text: str, fs_analysis: Optional[str] = None) -> str:
        """
        Generate a technical specification from a functional specification.
        
        Args:
            fs_text: The functional specification text
            fs_analysis: Optional analysis of the FS to guide TS generation
            
        Returns:
            Generated technical specification
        """
        if fs_analysis:
            self._add_thinking_step("Generating technical specification using both FS and analysis")
            return self._generate_with_analysis(fs_text, fs_analysis)
        else:
            self._add_thinking_step("Generating technical specification directly from FS")
            return self._generate_directly(fs_text)
    
    def _generate_directly(self, fs_text: str) -> str:
        """Generate TS directly from FS without analysis step."""
        prompt = """
Convert the following SAP Functional Specification into a detailed Technical Specification.
Include the following sections:
1. Executive Summary
2. SAP Objects List (tables, structures, programs, etc.)
3. Data Dictionary Changes
4. User Interface Changes
5. Process Flow
6. Technical Logic (with pseudocode or ABAP/UI5 code snippets where appropriate)
7. Error Handling
8. Authorization
9. Testing Strategy
10. Transport Strategy

Functional Specification:
"""
        
        messages = [
            {"role": "system", "content": "You are an SAP technical consultant specializing in creating technical specifications."},
            {"role": "user", "content": f"{prompt}\n\n{fs_text}"}
        ]
        
        return self._create_chat_completion(messages, temperature=0.3)
    
    def _generate_with_analysis(self, fs_text: str, fs_analysis: str) -> str:
        """Generate TS using both FS and analysis as input."""
        prompt = """
Convert the following SAP Functional Specification into a detailed Technical Specification.
Use the provided analysis to address any gaps and improve the quality of the technical specification.

Include the following sections:
1. Executive Summary
2. SAP Objects List (tables, structures, programs, etc.)
3. Data Dictionary Changes
4. User Interface Changes
5. Process Flow
6. Technical Logic (with pseudocode or ABAP/UI5 code snippets where appropriate)
7. Error Handling
8. Authorization
9. Testing Strategy
10. Transport Strategy

Functional Specification:
"""
        
        analysis_prompt = """
Analysis of the functional specification with identified gaps and recommendations:
"""
        
        messages = [
            {"role": "system", "content": "You are an SAP technical consultant specializing in creating detailed technical specifications."},
            {"role": "user", "content": f"{prompt}\n\n{fs_text}\n\n{analysis_prompt}\n\n{fs_analysis}"}
        ]
        
        return self._create_chat_completion(messages, temperature=0.3)
    
    def improve_ts(self, ts_content: str, feedback: str) -> str:
        """
        Improve a technical specification based on feedback.
        
        Args:
            ts_content: The current technical specification
            feedback: Feedback for improvement
            
        Returns:
            Improved technical specification
        """
        self._add_thinking_step("Improving technical specification based on feedback")
        
        prompt = """
Improve the following SAP Technical Specification based on the feedback provided:

Feedback:
"""
        
        messages = [
            {"role": "system", "content": "You are an SAP technical consultant specializing in creating detailed technical specifications."},
            {"role": "user", "content": f"{prompt}\n\n{feedback}\n\nCurrent Technical Specification:\n\n{ts_content}"}
        ]
        
        return self._create_chat_completion(messages, temperature=0.3)

###################
# Code Writer Agent
###################

class CodeWriterAgent(Agent):
    """
    Agent that generates SAP code from Technical Specifications.
    """
    
    def generate_code(self, ts_text: str, code_type: str = "ABAP") -> str:
        """
        Generate code from a technical specification.
        
        Args:
            ts_text: The technical specification text
            code_type: Type of code to generate ("ABAP" or "BTP")
            
        Returns:
            Generated code
        """
        self._add_thinking_step(f"Analyzing technical specification for {code_type} code generation")
        
        if code_type == "ABAP":
            self._add_thinking_step("Generating ABAP code")
            prompt = """
Generate complete, production-ready ABAP code based on the following Technical Specification.
Include appropriate error handling, comments, and follow ABAP best practices.
Focus on creating clean, maintainable, and efficient code.

Technical Specification:
"""
        else:
            self._add_thinking_step("Generating SAP BTP/Node.js code")
            prompt = """
Generate complete, production-ready SAP BTP/Node.js code based on the following Technical Specification.
Include appropriate error handling, comments, and follow JavaScript/Node.js best practices.
Focus on creating clean, maintainable, and efficient code structured for SAP Business Technology Platform.

Technical Specification:
"""
        
        result = self._generate_code(ts_text, prompt)
        return result
    
    def _generate_code(self, ts_text: str, prompt: str) -> str:
        """Generate code from TS with the specified prompt."""
        messages = [
            {"role": "system", "content": "You are an expert SAP developer specializing in creating well-structured code from technical specifications."},
            {"role": "user", "content": f"{prompt}\n\n{ts_text}"}
        ]
        
        return self._create_chat_completion(messages, temperature=0.2)
    
    def improve_code(self, code_content: str, feedback: str, code_type: str = "ABAP") -> str:
        """
        Improve code based on feedback.
        
        Args:
            code_content: The current code
            feedback: Feedback for improvement
            code_type: Type of code ("ABAP" or "BTP")
            
        Returns:
            Improved code
        """
        self._add_thinking_step(f"Improving {code_type} code based on feedback")
        
        if code_type == "ABAP":
            prompt = """
Improve the following ABAP code based on the feedback provided:

Feedback:
"""
        else:
            prompt = """
Improve the following SAP BTP/Node.js code based on the feedback provided:

Feedback:
"""
        
        messages = [
            {"role": "system", "content": "You are an expert SAP developer specializing in code optimization and improvement."},
            {"role": "user", "content": f"{prompt}\n\n{feedback}\n\nCurrent Code:\n\n{code_content}"}
        ]
        
        return self._create_chat_completion(messages, temperature=0.2)

###################
# Test Case Generator Agent
###################

class TestCaseGenerator(Agent):
    """
    Agent that generates unit tests from SAP code.
    """
    
    def generate_unit_tests(self, code: str, code_type: str = "ABAP") -> str:
        """
        Generate unit tests for the provided code.
        
        Args:
            code: The code to test
            code_type: Type of code ("ABAP" or "BTP"/"Node.js")
            
        Returns:
            Generated unit tests
        """
        self._add_thinking_step(f"Analyzing {code_type} code for unit test generation")
        
        if code_type == "ABAP":
            self._add_thinking_step("Generating ABAP unit tests")
            prompt = """
Generate comprehensive ABAP unit tests for the following code.
Use ABAP Unit framework with appropriate test classes, setup/teardown methods, and assertions.
Ensure good test coverage including positive cases, edge cases, and error scenarios.

ABAP Code to Test:
"""
        else:
            self._add_thinking_step("Generating JavaScript/Node.js unit tests")
            prompt = """
Generate comprehensive JavaScript unit tests for the following SAP BTP/Node.js code.
Use Jest testing framework with appropriate describe/it blocks, setup/teardown, and assertions.
Ensure good test coverage including positive cases, edge cases, and error scenarios.

JavaScript/Node.js Code to Test:
"""
        
        result = self._generate_tests(code, prompt)
        return result
    
    def _generate_tests(self, code: str, prompt: str) -> str:
        """Generate unit tests with the specified prompt."""
        messages = [
            {"role": "system", "content": "You are an expert SAP test engineer specializing in creating comprehensive unit tests."},
            {"role": "user", "content": f"{prompt}\n\n{code}"}
        ]
        
        return self._create_chat_completion(messages, temperature=0.3)
    
    def improve_tests(self, test_content: str, feedback: str, code_type: str = "ABAP") -> str:
        """
        Improve unit tests based on feedback.
        
        Args:
            test_content: The current unit tests
            feedback: Feedback for improvement
            code_type: Type of code ("ABAP" or "BTP")
            
        Returns:
            Improved unit tests
        """
        self._add_thinking_step(f"Improving {code_type} unit tests based on feedback")
        
        if code_type == "ABAP":
            prompt = """
Improve the following ABAP unit tests based on the feedback provided:

Feedback:
"""
        else:
            prompt = """
Improve the following JavaScript/Node.js unit tests based on the feedback provided:

Feedback:
"""
        
        messages = [
            {"role": "system", "content": "You are an expert SAP test engineer specializing in optimizing unit tests."},
            {"role": "user", "content": f"{prompt}\n\n{feedback}\n\nCurrent Unit Tests:\n\n{test_content}"}
        ]
        
        return self._create_chat_completion(messages, temperature=0.3)

###################
# Enhancement Agent
###################

class EnhancementSuggesterAgent(Agent):
    """
    Agent that suggests enhancements for SAP code.
    """
    
    def suggest_enhancements(self, code: str, code_type: str = "ABAP") -> str:
        """
        Suggest enhancements for the provided code.
        
        Args:
            code: The code to analyze
            code_type: Type of code ("ABAP" or "BTP")
            
        Returns:
            Enhancement suggestions
        """
        self._add_thinking_step(f"Analyzing {code_type} code for potential enhancements")
        
        if code_type == "ABAP":
            self._add_thinking_step("Generating ABAP enhancement suggestions")
            prompt = """
Analyze the following ABAP code and suggest comprehensive enhancements in the following categories:
1. Performance Optimization
2. Code Maintainability
3. Error Handling Improvements
4. Security Enhancements
5. Modern ABAP Features
6. Cloud Readiness (S/4HANA compatibility)

For each suggestion:
- Explain the issue or opportunity
- Provide concrete code examples for implementation
- Describe the benefits of the change

Format your response in markdown with clear sections.

ABAP Code to Analyze:
"""
        else:
            self._add_thinking_step("Generating BTP/Node.js enhancement suggestions")
            prompt = """
Analyze the following SAP BTP/Node.js code and suggest comprehensive enhancements in the following categories:
1. Performance Optimization
2. Code Maintainability
3. Error Handling Improvements
4. Security Enhancements
5. Modern JavaScript Features
6. Cloud-Native Improvements

For each suggestion:
- Explain the issue or opportunity
- Provide concrete code examples for implementation
- Describe the benefits of the change

Format your response in markdown with clear sections.

BTP/Node.js Code to Analyze:
"""
        
        result = self._generate_enhancements(code, code_type)
        return result
    
    def _generate_enhancements(self, code: str, code_type: str) -> str:
        """Generate enhancement suggestions."""
        if code_type == "ABAP":
            prompt = """
Analyze the following ABAP code and suggest comprehensive enhancements in the following categories:
1. Performance Optimization
2. Code Maintainability
3. Error Handling Improvements
4. Security Enhancements
5. Modern ABAP Features
6. Cloud Readiness (S/4HANA compatibility)

For each suggestion:
- Explain the issue or opportunity
- Provide concrete code examples for implementation
- Describe the benefits of the change

Format your response in markdown with clear sections.

ABAP Code to Analyze:
"""
        else:
            prompt = """
Analyze the following SAP BTP/Node.js code and suggest comprehensive enhancements in the following categories:
1. Performance Optimization
2. Code Maintainability
3. Error Handling Improvements
4. Security Enhancements
5. Modern JavaScript Features
6. Cloud-Native Improvements

For each suggestion:
- Explain the issue or opportunity
- Provide concrete code examples for implementation
- Describe the benefits of the change

Format your response in markdown with clear sections.

BTP/Node.js Code to Analyze:
"""
        
        messages = [
            {"role": "system", "content": "You are an expert SAP developer specializing in code optimization and modernization."},
            {"role": "user", "content": f"{prompt}\n\n{code}"}
        ]
        
        return self._create_chat_completion(messages, temperature=0.4)

###################
# Transport Doc Agent
###################

class TransportDocGenerator(Agent):
    """
    Agent that generates SAP transport request documentation.
    """
    
    def generate_transport_doc(self, content: str, content_type: str = "CODE") -> str:
        """
        Generate transport request documentation.
        
        Args:
            content: The code or specification to document
            content_type: Type of content ("CODE", "TS", or "FS")
            
        Returns:
            Generated transport documentation
        """
        self._add_thinking_step(f"Analyzing {content_type} to generate transport request documentation")
        result = self._generate_doc(content, content_type)
        return result
    
    def _generate_doc(self, content: str, content_type: str) -> str:
        """Generate transport documentation."""
        prompt_intro = "Generate comprehensive SAP transport request documentation based on the following "
        
        if content_type == "CODE":
            prompt = f"{prompt_intro}code. Include sections for objects transported, purpose, impact analysis, testing recommendations, prerequisite transports, and post-implementation steps."
        elif content_type == "TS":
            prompt = f"{prompt_intro}technical specification. Focus on the technical objects that will need to be transported, dependencies, sequence, and verification steps."
        else:  # FS
            prompt = f"{prompt_intro}functional specification. Extract the key SAP objects likely to be affected, and create transport documentation focusing on business impact, risk assessment, and testing requirements."
        
        messages = [
            {"role": "system", "content": "You are an expert SAP transport manager specializing in creating comprehensive transport documentation."},
            {"role": "user", "content": f"{prompt}\n\n{content}"}
        ]
        
        return self._create_chat_completion(messages, temperature=0.3)

###################
# Multi-Agent Orchestrator
###################

###################
# AutoGen Integration
###################

class AutoGenIntegration:
    """
    Integration with Microsoft's AutoGen framework for agent-to-agent collaboration,
    using Groq as the LLM provider.
    """
    
    def __init__(self, groq_client: GroqClient = None):
        """
        Initialize the AutoGen integration with a Groq client.
        
        Args:
            groq_client: The Groq client for LLM integration
        """
        if groq_client is None:
            groq_client = GroqClient()
        
        self.groq_client = groq_client
        self.agents = {}
        
        # Configure OpenAI client with Groq API
        openai.api_key = self.groq_client.api_key
        openai.base_url = "https://api.groq.com/openai/v1"
        
        # Create config list for AutoGen
        self.config_list = self._create_config_list()
        
    def _create_config_list(self) -> List[Dict[str, Any]]:
        """
        Create the configuration list for AutoGen agents using Groq.
        
        Returns:
            List of configuration dictionaries for AutoGen
        """
        # AutoGen requires 'gpt-4' format names, so we'll map Groq models to compatible format
        model_mapping = {
            "llama-3.3-70b-versatile": "llama-3.3-70b-versatile",
            "llama3-8b-8192": "llama-3.3-70b-versatileB",
            "mixtral-8x7b-32768": "llama-3.3-70b-versatile",
            # Add more mappings as needed
        }
        
        # Use gpt-4 as default fallback
        mapped_model = model_mapping.get(self.groq_client.model, "llama-3.3-70b-versatile")
        
        return [
            {
                "model": mapped_model,  # Using compatible model name
                "api_key": self.groq_client.api_key,
                "base_url": "https://api.groq.com/openai/v1",
                "api_type": "openai"
            }
        ]
    
    def create_agents(self, task_type: str) -> None:
        """
        Create the appropriate agents based on the task type.
        
        Args:
            task_type: Type of task ("fs_to_ts", "ts_to_code", "code_to_test", etc.)
        """
        # Define the configurations for different agent types
        
        config_list1 = [{"model": "llama-3.3-70b-versatile", "api_key": api_key, "api_type": "groq"}]

        llm_config = {
                        "temperature": 0,
                        "config_list": config_list1,
                    }

     
        
        
        # Create the user proxy agent (interface between user and system)
        self.agents["user_proxy"] = autogen.UserProxyAgent(
            name="User",
            human_input_mode="NEVER",
            max_consecutive_auto_reply=0,
            code_execution_config={"use_docker": False},
            system_message="You are a user who needs help with an SAP-related task."
        )
        
        # Create specialized agents based on task type
        if task_type == "fs_to_ts":
            # For FS to TS conversion, create an analyst and a technical writer
            self.agents["analyst"] = autogen.AssistantAgent(
                name="SAP_Analyst",
                system_message="""You are an expert SAP business analyst. Your job is to analyze functional specifications 
                and extract the key business requirements and needs. Focus on understanding the business process, 
                data requirements, and integration points.""",
                llm_config=llm_config
            )
            
            self.agents["tech_writer"] = autogen.AssistantAgent(
                name="Technical_Writer",
                system_message="""You are a technical specification writer for SAP systems. Your task is to create detailed
                technical specifications from functional requirements. Include all necessary SAP objects, data structures,
                interface definitions, and technical implementation details.""",
                llm_config=llm_config
            )
            
        elif task_type == "ts_to_code":
            # For TS to Code conversion, create an architect and a developer
            self.agents["architect"] = autogen.AssistantAgent(
                name="SAP_Architect",
                system_message="""You are an expert SAP solution architect. Your job is to design the technical solution
                based on a technical specification. Define the technical approach, components, and overall structure.""",
                llm_config=llm_config
            )
            
            self.agents["developer"] = autogen.AssistantAgent(
                name="SAP_Developer",
                system_message="""You are an expert SAP developer with deep knowledge of ABAP and BTP development.
                Your task is to write clean, efficient, and maintainable code that implements the technical specification
                and follows SAP best practices.""",
                llm_config=llm_config
            )
            
        elif task_type == "code_to_test":
            # For Code to Test conversion, create a test engineer and a quality analyst
            self.agents["test_engineer"] = autogen.AssistantAgent(
                name="Test_Engineer",
                system_message="""You are an expert SAP test engineer. Your job is to write comprehensive unit tests
                for SAP code. Focus on test coverage, edge cases, and proper test isolation.""",
                llm_config=llm_config
            )
            
            self.agents["qa_analyst"] = autogen.AssistantAgent(
                name="QA_Analyst",
                system_message="""You are a quality assurance specialist for SAP solutions. Your task is to review
                and enhance test cases to ensure they properly validate the code and cover all required scenarios.""",
                llm_config=llm_config
            )
            
        elif task_type == "enhancement":
            # For enhancement suggestions, create a code reviewer and an optimization expert
            self.agents["code_reviewer"] = autogen.AssistantAgent(
                name="Code_Reviewer",
                system_message="""You are an expert SAP code reviewer. Your job is to analyze code for potential
                improvements, identify any issues, and suggest best practices. Focus on readability and maintainability.""",
                llm_config=llm_config
            )
            
            self.agents["optimization_expert"] = autogen.AssistantAgent(
                name="Optimization_Expert",
                system_message="""You are an SAP performance optimization expert. Your task is to identify performance
                bottlenecks and suggest improvements. Focus on database access, memory usage, and algorithmic efficiency.""",
                llm_config=llm_config
            )
            
        elif task_type == "transport_docs":
            # For transport documentation, create a documentation specialist and a transport coordinator
            self.agents["doc_specialist"] = autogen.AssistantAgent(
                name="Documentation_Specialist",
                system_message="""You are an SAP documentation specialist. Your job is to create comprehensive
                and clear documentation for transport requests. Include all affected objects, dependencies, and implementation steps.""",
                llm_config=llm_config
            )
            
            self.agents["transport_coordinator"] = autogen.AssistantAgent(
                name="Transport_Coordinator",
                system_message="""You are an SAP transport coordinator. Your task is to ensure transport requests are
                properly documented and include all relevant information for successful transport between systems.""",
                llm_config=llm_config
            )
        
        # Create a manager agent for all scenarios to coordinate the conversation
        self.agents["manager"] = autogen.AssistantAgent(
            name="SAP_Manager",
            system_message="""You are a project manager for SAP implementations. Your role is to coordinate the 
            conversation between different experts, ensure the final output meets requirements, and summarize the results.
            You have the final say on the output format and content.""",
            llm_config=llm_config
        )
    
    def execute_task(self, task_description: str, content: str, task_type: str) -> str:
        """
        Execute a task using AutoGen agents with agent-to-agent collaboration.
        
        Args:
            task_description: Description of the task to execute
            content: Content information (code, specification, etc.)
            task_type: Type of task (fs_to_ts, ts_to_code, etc.)
            
        Returns:
            Result of the agent-to-agent collaboration
        """
        # Create agents based on task type
        self.create_agents(task_type)
        
        # Set up the AutoGen group chat
        groupchat = autogen.GroupChat(
            agents=list(self.agents.values()),
            messages=[],
            max_round=10
        )
        
        # Create a group chat manager
        manager = autogen.GroupChatManager(
            groupchat=groupchat,
            llm_config={"config_list": self.config_list}
        )
        
        # Prepare the task message
        task_message = f"""
        # Task: {task_description}
        
        # Content:
        {content}
        
        Please collaborate to complete this task. Each expert should contribute their perspective.
        The final output should be a comprehensive solution that meets all requirements.
        """
        
        # Initiate the chat with the task message
        self.agents["user_proxy"].initiate_chat(
            manager,
            message=task_message
        )
        
        # Extract the final result from the chat history
        chat_history = groupchat.messages
        
        # Get the last message from the SAP_Manager as the final result
        final_messages = [msg for msg in chat_history if msg.get("sender") == "SAP_Manager"]
        
        if final_messages:
            # Get the most recent message from the manager
            final_result = final_messages[-1].get("content", "No result generated")
            
            # Clean up the result
            final_result = self._clean_result(final_result)
            
            return final_result
        else:
            # Fallback if no manager message is found
            return "No clear result was generated from the agent collaboration. Please try again with more specific instructions."
    
    def _clean_result(self, result: str) -> str:
        """
        Clean up the result text to remove conversation artifacts.
        
        Args:
            result: The raw result text from agent conversation
            
        Returns:
            Cleaned result text
        """
        # Remove any lines that reference the conversation or other agents
        clean_lines = []
        for line in result.split("\n"):
            if not any(agent_name in line for agent_name in self.agents.keys() if agent_name != "manager"):
                if not any(phrase in line.lower() for phrase in ["i agree", "thank you", "let me", "i'll", "i'd", "i've"]):
                    clean_lines.append(line)
                    
        return "\n".join(clean_lines)

###################
# Multi-Agent Orchestrator
###################

class MultiAgentOrchestrator(Agent):
    """
    Orchestrator for multi-agent collaboration.
    Manages communication between specialized SAP agents.
    """
    
    def __init__(self, groq_client: GroqClient):
        super().__init__(groq_client)
        self.fs_analyzer = FSAnalyzerAgent(groq_client)
        self.ts_generator = TSGeneratorAgent(groq_client)
        self.code_writer = CodeWriterAgent(groq_client)
        self.test_generator = TestCaseGenerator(groq_client)
        self.enhancement_suggester = EnhancementSuggesterAgent(groq_client)
        self.transport_doc_generator = TransportDocGenerator(groq_client)
    
    def execute_task(self, task_description: str, context: str) -> str:
        """
        Execute a task with multi-agent collaboration.
        
        Args:
            task_description: Description of the task to execute
            context: Context information (code, specification, etc.)
            
        Returns:
            Result of the multi-agent execution
        """
        # Plan the execution by determining which agents should be involved
        self._add_thinking_step("Planning multi-agent task execution")
        
        # Simple keyword-based task routing
        task_lower = task_description.lower()
        
        if "technical specification" in task_lower and "functional specification" in task_lower:
            # FS to TS flow
            self._add_thinking_step("Routing task to FS Analyzer and TS Generator agents")
            fs_analysis = self.fs_analyzer.analyze_fs(context)
            self.thinking_steps.extend(self.fs_analyzer.get_thinking_steps())
            
            result = self.ts_generator.generate_ts(context, fs_analysis)
            self.thinking_steps.extend(self.ts_generator.get_thinking_steps())
            
        elif "code" in task_lower and "technical specification" in task_lower:
            # TS to Code flow
            self._add_thinking_step("Routing task to Code Writer agent")
            code_type = "ABAP"
            if "btp" in task_lower or "node" in task_lower:
                code_type = "BTP"
                
            result = self.code_writer.generate_code(context, code_type)
            self.thinking_steps.extend(self.code_writer.get_thinking_steps())
            
        elif "test" in task_lower and "code" in task_lower:
            # Code to Test flow
            self._add_thinking_step("Routing task to Test Case Generator agent")
            code_type = "ABAP"
            if "btp" in task_lower or "node" in task_lower:
                code_type = "Node.js"
                
            result = self.test_generator.generate_unit_tests(context, code_type)
            self.thinking_steps.extend(self.test_generator.get_thinking_steps())
            
        elif "enhance" in task_lower or "suggestion" in task_lower:
            # Enhancement suggestions flow
            self._add_thinking_step("Routing task to Enhancement Suggester agent")
            code_type = "ABAP"
            if "btp" in task_lower or "node" in task_lower:
                code_type = "BTP"
                
            result = self.enhancement_suggester.suggest_enhancements(context, code_type)
            self.thinking_steps.extend(self.enhancement_suggester.get_thinking_steps())
            
        elif "transport" in task_lower or "documentation" in task_lower:
            # Transport documentation flow
            self._add_thinking_step("Routing task to Transport Doc Generator agent")
            content_type = "CODE"
            if "technical specification" in task_lower:
                content_type = "TS"
            elif "functional specification" in task_lower:
                content_type = "FS"
                
            result = self.transport_doc_generator.generate_transport_doc(context, content_type)
            self.thinking_steps.extend(self.transport_doc_generator.get_thinking_steps())
            
        else:
            # Fallback to general execution via simulated conversation
            self._add_thinking_step("Using generic multi-agent conversation for unspecified task type")
            result = self._simulate_agent_conversation(task_description, context)
        
        return result
    
    def _simulate_agent_conversation(self, task_description: str, context: str) -> str:
        """Simulate a conversation between agents to solve a general task."""
        self._add_thinking_step("Starting simulated conversation between agents")
        
        prompt = f"""
Task: {task_description}

Context:
{context}

Please simulate a conversation between specialized SAP agents to solve this task collaboratively:
1. Analyst Agent - Analyzes requirements and plans the approach
2. Technical Designer - Creates technical design solutions
3. Developer - Implements code and technical solutions
4. Quality Assurance - Tests and validates the solution
5. Documentation Specialist - Creates documentation and explanations

The agents should discuss the task and collaborate to produce a comprehensive solution.
"""
        
        messages = [
            {"role": "system", "content": "You are an multi-agent orchestrator simulating a conversation between specialized SAP agents."},
            {"role": "user", "content": prompt}
        ]
        
        conversation = self._create_chat_completion(messages, temperature=0.5)
        
        # Extract the final result/solution from the conversation
        result_prompt = f"""
Based on the following simulated conversation between SAP expert agents, extract and format the final solution that addresses this task:

Task: {task_description}

Conversation:
{conversation}

Format the solution in a clean, well-structured way without any conversation elements.
"""
        
        messages = [
            {"role": "system", "content": "You are a solution extractor that formats and presents the final solution from multi-agent conversations."},
            {"role": "user", "content": result_prompt}
        ]
        
        final_result = self._create_chat_completion(messages, temperature=0.3)
        self._add_thinking_step("Extracted final solution from agent conversation")
        
        return final_result

###################
# Agent Factory
###################

class AgentFactory:
    """
    Factory for creating different types of agents with consistent configuration.
    """
    
    @staticmethod
    def create_fs_analyzer_agent(groq_client: GroqClient = None):
        """Create an agent for analyzing functional specifications."""
        if groq_client is None:
            groq_client = GroqClient()
        return FSAnalyzerAgent(groq_client)
    
    @staticmethod
    def create_ts_generator_agent(groq_client: GroqClient = None):
        """Create an agent for generating technical specifications."""
        if groq_client is None:
            groq_client = GroqClient()
        return TSGeneratorAgent(groq_client)
    
    @staticmethod
    def create_code_writer_agent(groq_client: GroqClient = None):
        """Create an agent for generating code."""
        if groq_client is None:
            groq_client = GroqClient()
        return CodeWriterAgent(groq_client)
    
    @staticmethod
    def create_test_case_generator(groq_client: GroqClient = None):
        """Create an agent for generating unit tests."""
        if groq_client is None:
            groq_client = GroqClient()
        return TestCaseGenerator(groq_client)
    
    @staticmethod
    def create_enhancement_suggester_agent(groq_client: GroqClient = None):
        """Create an agent for suggesting enhancements."""
        if groq_client is None:
            groq_client = GroqClient()
        return EnhancementSuggesterAgent(groq_client)
    
    @staticmethod
    def create_transport_doc_generator(groq_client: GroqClient = None):
        """Create an agent for generating transport documentation."""
        if groq_client is None:
            groq_client = GroqClient()
        return TransportDocGenerator(groq_client)
        
    @staticmethod
    def create_autogen_integration(groq_client: GroqClient = None):
        """Create an AutoGen integration for agent-to-agent collaboration."""
        if groq_client is None:
            groq_client = GroqClient()
        return AutoGenIntegration(groq_client)

###################
# Page Functions
###################

# FS to TS Page Function
def show_fs_to_ts():
    """Display the FS to TS automation page."""
    st.header("FS → TS Automation")
    st.markdown("""
        Convert Functional Specifications into detailed Technical Specifications. 
        Upload or paste your Functional Specification text below.
    """)
    
    # Initialize session state for this page
    if 'fs_content' not in st.session_state:
        st.session_state.fs_content = ""
    if 'ts_result' not in st.session_state:
        st.session_state.ts_result = ""
    if 'fs_analysis' not in st.session_state:
        st.session_state.fs_analysis = ""
    if 'show_thinking' not in st.session_state:
        st.session_state.show_thinking = False
    if 'thinking_steps' not in st.session_state:
        st.session_state.thinking_steps = []
    if 'use_agent_framework' not in st.session_state:
        st.session_state.use_agent_framework = False
    if 'use_autogen' not in st.session_state:
        st.session_state.use_autogen = False
    
    # File upload or text input
    st.subheader("Input Functional Specification")
    
    tab1, tab2 = st.tabs(["Upload File", "Paste Text"])
    
    with tab1:
        uploaded_file = st.file_uploader("Upload a Functional Specification document", 
                                       type=['txt', 'pdf', 'docx'],
                                       help="Supported formats: .txt, .pdf, .docx")
        
        if uploaded_file is not None:
            try:
                # Extract text from the uploaded file
                fs_text = DocumentProcessor.extract_text(uploaded_file)
                st.session_state.fs_content = fs_text
                st.success(f"Successfully extracted text from {uploaded_file.name}")
                
                # Show preview
                with st.expander("Preview extracted text"):
                    st.text_area("File content", fs_text, height=200)
            except Exception as e:
                st.error(f"Error processing file: {str(e)}")
    
    with tab2:
        fs_text_input = st.text_area("Enter Functional Specification text", 
                                   height=200,
                                   value=st.session_state.fs_content)
        
        if fs_text_input != st.session_state.fs_content:
            st.session_state.fs_content = fs_text_input
    
    # Options
    st.subheader("Options")
    
    model_options = ["llama-3.3-70b-versatile", "llama3-8b-8192", "mixtral-8x7b-32768","deepseek-r1-distill-llama-70b"]
    selected_model = st.selectbox("Select LLM Model", model_options)
    
    col1, col2 = st.columns(2)
    
    with col1:
        analyze_first = st.checkbox("Analyze FS before generating TS", value=True,
                                  help="Perform analysis of the FS to improve TS quality")
    
    with col2:
        st.session_state.use_agent_framework = st.checkbox("Use Multi-Agent Framework", value=False,
                                       help="Use multiple specialized agents to generate the TS")
    
    # Advanced options
    with st.expander("Advanced Options"):
        st.session_state.show_thinking = st.checkbox("Show agent thinking process", value=False)
        st.session_state.use_autogen = st.checkbox("Use Microsoft AutoGen Framework", value=False, help="Enables agent-to-agent collaboration using Microsoft AutoGen framework")
    
    # Generate button
    if st.button("Generate Technical Specification", type="primary", 
               disabled=not st.session_state.fs_content):
        
        if not st.session_state.fs_content.strip():
            st.error("Please provide Functional Specification content first.")
            return
        
        with st.spinner("Generating Technical Specification..."):
            try:
                # Initialize GROQ client with selected model
                groq_client = GroqClient()
                groq_client.set_model(selected_model)
                
                if st.session_state.use_autogen:
                    # Use AutoGen integration
                    autogen_integration = AgentFactory.create_autogen_integration(groq_client)
                    
                    # Execute the task
                    result = autogen_integration.execute_task(
                        "Convert this functional specification into a detailed SAP technical specification.",
                        st.session_state.fs_content,
                        "fs_to_ts"
                    )
                    
                    st.session_state.ts_result = result
                    st.session_state.thinking_steps = ["Using Microsoft AutoGen agent collaboration"]
                    
                elif st.session_state.use_agent_framework:
                    # Use multi-agent framework
                    orchestrator = MultiAgentOrchestrator(groq_client)
                    
                    # Execute the task
                    result = orchestrator.execute_task(
                        "Convert this functional specification into a detailed SAP technical specification.",
                        st.session_state.fs_content
                    )
                    
                    st.session_state.ts_result = result
                    st.session_state.thinking_steps = orchestrator.get_thinking_steps()
                    
                else:
                    # Use single agents in sequence
                    if analyze_first:
                        # First analyze the FS
                        fs_analyzer = AgentFactory.create_fs_analyzer_agent(groq_client)
                        st.session_state.fs_analysis = fs_analyzer.analyze_fs(st.session_state.fs_content)
                        st.session_state.thinking_steps = fs_analyzer.get_thinking_steps()
                        
                        # Then generate TS based on FS and analysis
                        ts_generator = AgentFactory.create_ts_generator_agent(groq_client)
                        st.session_state.ts_result = ts_generator.generate_ts(
                            st.session_state.fs_content, 
                            st.session_state.fs_analysis
                        )
                        
                        # Combine thinking steps
                        st.session_state.thinking_steps.extend(ts_generator.get_thinking_steps())
                        
                    else:
                        # Generate TS directly from FS
                        ts_generator = AgentFactory.create_ts_generator_agent(groq_client)
                        st.session_state.ts_result = ts_generator.generate_ts(st.session_state.fs_content)
                        st.session_state.thinking_steps = ts_generator.get_thinking_steps()
                
                st.success("Technical Specification generated successfully!")
                
            except Exception as e:
                st.error(f"Error generating Technical Specification: {str(e)}")
    
    # Display results if available
    if st.session_state.ts_result:
        st.subheader("Generated Technical Specification")
        
        # Display thinking process if enabled
        if st.session_state.show_thinking and st.session_state.thinking_steps:
            with st.expander("Agent Thinking Process", expanded=False):
                for i, step in enumerate(st.session_state.thinking_steps):
                    st.markdown(f"{i+1}. {step}")
        
        # Display analysis if available
        if st.session_state.fs_analysis:
            with st.expander("Functional Specification Analysis", expanded=False):
                st.markdown(st.session_state.fs_analysis)
        
        # Display the generated TS
        st.markdown(st.session_state.ts_result)
        
        # Download options
        st.subheader("Download Options")
        
        col1, col2 = st.columns(2)
        
        with col1:
            # Download as text
            FileHandler.display_download_button(
                st.session_state.ts_result,
                "technical_specification.txt",
                "Download as Text (.txt)",
                "text/plain"
            )
        
        with col2:
            # Download as Word document
            docx_bytes = FileHandler.create_docx_download(st.session_state.ts_result)
            FileHandler.display_download_button(
                docx_bytes,
                "technical_specification.docx",
                "Download as Word (.docx)",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    
    # Sample example
    with st.expander("See Example", expanded=False):
        st.markdown("### Sample Functional Specification")
        sample_fs = """
        # Functional Specification: Custom Field Extension for SAP Sales Order Creation (VA01)
        
        ## 1. Overview
        This specification outlines the requirement to add a custom field "ZDELIVERY_PRIORITY" to the SAP Sales Order Creation transaction (VA01) to capture delivery priority information for special customer orders.
        
        ## 2. Business Requirement
        The sales team needs to indicate special delivery priorities for certain customers. This information needs to be visible during order creation and should be carried through to subsequent delivery documents.
        
        ## 3. Field Details
        - Field Name: ZDELIVERY_PRIORITY
        - Data Type: Character
        - Length: 2
        - Values:
          - 'H1': Highest Priority (Same Day)
          - 'H2': High Priority (Next Day)
          - 'M1': Medium Priority (2-3 Days)
          - 'L1': Low Priority (Standard)
        - Default Value: 'L1'
        
        ## 4. UI Placement
        The field should be placed in the Sales tab of the VA01 transaction, directly below the Delivery Priority field.
        
        ## 5. Business Process Impact
        - This field will be used by the warehouse and shipping teams to prioritize picking and packing.
        - Reports will be created to sort orders by this priority field.
        - The field should flow through to delivery documents and be visible in VL01N.
        
        ## 6. Authorizations
        All users with access to VA01 should be able to view and modify this field.
        
        ## 7. Approval
        Project Manager: John Smith
        Business Owner: Sarah Johnson
        Date: 2023-06-15
        """
        
        st.text_area("Example Functional Specification", sample_fs, height=200)
        
        if st.button("Use Example"):
            st.session_state.fs_content = sample_fs
            st.rerun()

# TS to Code Page Function
def show_ts_to_code():
    """Display the TS to Code automation page."""
    st.header("TS → Code Automation")
    st.markdown("""
        Generate SAP code from Technical Specifications. 
        Upload or paste your Technical Specification text below.
    """)
    
    # Initialize session state for this page
    if 'ts_content' not in st.session_state:
        st.session_state.ts_content = ""
    if 'code_result' not in st.session_state:
        st.session_state.code_result = ""
    if 'show_thinking' not in st.session_state:
        st.session_state.show_thinking = False
    if 'thinking_steps' not in st.session_state:
        st.session_state.thinking_steps = []
    if 'use_agent_framework' not in st.session_state:
        st.session_state.use_agent_framework = False
    
    # File upload or text input
    st.subheader("Input Technical Specification")
    
    tab1, tab2 = st.tabs(["Upload File", "Paste Text"])
    
    with tab1:
        uploaded_file = st.file_uploader("Upload a Technical Specification document", 
                                        type=['txt', 'pdf', 'docx'],
                                        help="Supported formats: .txt, .pdf, .docx")
        
        if uploaded_file is not None:
            try:
                # Extract text from the uploaded file
                ts_text = DocumentProcessor.extract_text(uploaded_file)
                st.session_state.ts_content = ts_text
                st.success(f"Successfully extracted text from {uploaded_file.name}")
                
                # Show preview
                with st.expander("Preview extracted text"):
                    st.text_area("File content", ts_text, height=200)
            except Exception as e:
                st.error(f"Error processing file: {str(e)}")
    
    with tab2:
        ts_text_input = st.text_area("Enter Technical Specification text", 
                                    height=200,
                                    value=st.session_state.ts_content)
        
        if ts_text_input != st.session_state.ts_content:
            st.session_state.ts_content = ts_text_input
    
    # Options
    st.subheader("Options")
    
    model_options = ["llama-3.3-70b-versatile", "llama3-8b-8192", "mixtral-8x7b-32768","deepseek-r1-distill-llama-70b"]
    selected_model = st.selectbox("Select LLM Model", model_options)
    
    col1, col2 = st.columns(2)
    
    with col1:
        code_type = st.radio("Code Type", ["ABAP", "BTP/Node.js"], 
                           help="Select the type of code to generate")
    
    with col2:
        st.session_state.use_agent_framework = st.checkbox("Use Multi-Agent Framework", value=False,
                                        help="Use multiple specialized agents to generate the code")
    
    # Advanced options
    with st.expander("Advanced Options"):
        st.session_state.use_autogen = st.checkbox("Use Microsoft AutoGen Framework", value=False, help="Enables agent-to-agent collaboration using Microsoft AutoGen framework")
        st.session_state.show_thinking = st.checkbox("Show agent thinking process", value=False)
    
    # Generate button
    if st.button("Generate Code", type="primary", 
                disabled=not st.session_state.ts_content):
        
        if not st.session_state.ts_content.strip():
            st.error("Please provide Technical Specification content first.")
            return
        
        with st.spinner("Generating Code..."):
            try:
                # Initialize GROQ client with selected model
                groq_client = GroqClient()
                groq_client.set_model(selected_model)
                
                # Determine code type for processing
                code_type_value = "ABAP" if code_type == "ABAP" else "BTP"
                
                if st.session_state.use_autogen:
                    # Use AutoGen integration
                    autogen_integration = AgentFactory.create_autogen_integration(groq_client)
                    
                    # Execute the task
                    result = autogen_integration.execute_task(
                        f"Convert this technical specification into {code_type_value} code.",
                        st.session_state.ts_content,
                        "ts_to_code"
                    )
                    
                    st.session_state.code_result = result
                    st.session_state.thinking_steps = ["Using Microsoft AutoGen agent collaboration"]
                    
                elif st.session_state.use_agent_framework:
                    # Use multi-agent framework
                    orchestrator = MultiAgentOrchestrator(groq_client)
                    
                    # Execute the task
                    result = orchestrator.execute_task(
                        f"Convert this technical specification into {code_type_value} code.",
                        st.session_state.ts_content
                    )
                    
                    st.session_state.code_result = result
                    st.session_state.thinking_steps = orchestrator.get_thinking_steps()
                    
                else:
                    # Use single agent
                    code_writer = AgentFactory.create_code_writer_agent(groq_client)
                    st.session_state.code_result = code_writer.generate_code(
                        st.session_state.ts_content,
                        code_type_value
                    )
                    st.session_state.thinking_steps = code_writer.get_thinking_steps()
                
                st.success("Code generated successfully!")
                
            except Exception as e:
                st.error(f"Error generating code: {str(e)}")
    
    # Display results if available
    if st.session_state.code_result:
        st.subheader("Generated Code")
        
        # Display thinking process if enabled
        if st.session_state.show_thinking and st.session_state.thinking_steps:
            with st.expander("Agent Thinking Process", expanded=False):
                for i, step in enumerate(st.session_state.thinking_steps):
                    st.markdown(f"{i+1}. {step}")
        
        # Display the generated code
        st.code(st.session_state.code_result)
        
        # Feedback and regeneration
        feedback = st.text_area("Feedback for improvement (optional)")
        
        if feedback and st.button("Improve Code Based on Feedback"):
            with st.spinner("Improving code..."):
                try:
                    # Initialize GROQ client with selected model
                    groq_client = GroqClient()
                    groq_client.set_model(selected_model)
                    
                    # Determine code type for processing
                    code_type_value = "ABAP" if code_type == "ABAP" else "BTP"
                    
                    # Use code writer to improve the code
                    code_writer = AgentFactory.create_code_writer_agent(groq_client)
                    improved_code = code_writer.improve_code(
                        st.session_state.code_result,
                        feedback,
                        code_type_value
                    )
                    
                    # Update the result
                    st.session_state.code_result = improved_code
                    st.session_state.thinking_steps.extend(code_writer.get_thinking_steps())
                    
                    st.success("Code improved successfully!")
                    st.rerun()
                    
                except Exception as e:
                    st.error(f"Error improving code: {str(e)}")
        
        # Download options
        st.subheader("Download Options")
        
        code_ext = ".abap" if code_type == "ABAP" else ".js"
        
        FileHandler.display_download_button(
            st.session_state.code_result,
            f"generated_code{code_ext}",
            f"Download Code ({code_ext})",
            "text/plain"
        )
        
        # Option to create package with multiple files
        with st.expander("Create Code Package"):
            st.markdown("Create a zip package containing all generated code files.")
            
            zip_files = {}
            
            if code_type == "ABAP":
                # Add files for ABAP code
                zip_files[f"zcustom_program{code_ext}"] = st.session_state.code_result
                zip_files["README.md"] = f"# Generated ABAP Code\n\nThis package contains ABAP code generated from a Technical Specification.\n\n## Implementation Notes\n\nPlease review the code before implementing in your SAP system."
            else:
                # Add files for BTP/Node.js code
                zip_files["index.js"] = st.session_state.code_result
                zip_files["README.md"] = f"# Generated SAP BTP Code\n\nThis package contains SAP BTP/Node.js code generated from a Technical Specification.\n\n## Implementation Notes\n\nPlease review the code before implementing in your SAP BTP environment."
            
            zip_bytes = FileHandler.create_zip_download(zip_files)
            
            FileHandler.display_download_button(
                zip_bytes,
                "sap_code_package.zip",
                "Download Code Package (.zip)",
                "application/zip"
            )
    
    # Sample example
    with st.expander("See Example", expanded=False):
        st.markdown("### Sample Technical Specification")
        sample_ts = """
        # Technical Specification: Custom Field Extension for SAP Sales Order Creation (VA01)
        
        ## 1. Executive Summary
        This technical specification details the implementation of a new custom field "ZDELIVERY_PRIORITY" in the Sales Order transaction (VA01). The field will capture special delivery priority instructions that will flow through to delivery documents.
        
        ## 2. SAP Objects List
        - Custom Data Element: ZDE_DELIVERY_PRIORITY
        - Custom Domain: ZDO_DELIVERY_PRIORITY
        - Append Structure: ZVBAK_APPEND (for table VBAK)
        - Custom Table: ZDELIVERY_PRIORITY_TEXT (for priority descriptions)
        - Transaction Enhancement: VA01 screen modification
        - Custom Program: ZDELIVERY_PRIORITY_REPORT (for reporting)
        
        ## 3. Data Dictionary Changes
        ### 3.1 Domain: ZDO_DELIVERY_PRIORITY
        - Data Type: CHAR
        - Length: 2
        - Value Table: ZDELIVERY_PRIORITY_TEXT
        
        ### 3.2 Data Element: ZDE_DELIVERY_PRIORITY
        - Domain: ZDO_DELIVERY_PRIORITY
        - Short Description: "Delivery Priority Code"
        - Field Labels:
          - Short: "Del. Priority"
          - Medium: "Deliv.Priority"
          - Long: "Delivery Priority"
          - Heading: "Delivery Priority Code"
        
        ### 3.3 Append Structure: ZVBAK_APPEND
        - Field: ZDELIVERY_PRIORITY
        - Data Element: ZDE_DELIVERY_PRIORITY
        
        ### 3.4 Custom Table: ZDELIVERY_PRIORITY_TEXT
        - Fields:
          - PRIORITY_CODE: CHAR(2) - Primary Key
          - LANGUAGE: LANG - Primary Key
          - DESCRIPTION: CHAR(40)
        - Initial Data:
          - 'H1', 'E', 'Highest Priority (Same Day)'
          - 'H2', 'E', 'High Priority (Next Day)'
          - 'M1', 'E', 'Medium Priority (2-3 Days)'
          - 'L1', 'E', 'Low Priority (Standard)'
        
        ## 4. User Interface Changes
        ### 4.1 VA01 Screen Modification
        - Screen: SAPMV45A, SUBSCREEN 0101
        - Add field ZDELIVERY_PRIORITY to Sales tab
        - Position: Below standard Delivery Priority field
        - Add F4 help functionality linked to ZDELIVERY_PRIORITY_TEXT
        
        ## 5. Process Flow
        1. User opens VA01 to create a new sales order
        2. User enters or selects a value for ZDELIVERY_PRIORITY
        3. Value is stored in VBAK-ZDELIVERY_PRIORITY via append structure
        4. Field is carried over to delivery document (LIKP) via user exit
        
        ## 6. Technical Logic
        ### 6.1 User Exit for Field Transport
        - Implement user exit USEREXIT_DELIVERYHEADER_PREPARE
        - Copy VBAK-ZDELIVERY_PRIORITY to LIKP-ZDELIVERY_PRIORITY
        
        ### 6.2 Field Default Logic
        - Implement BAdI BADI_SD_SALESORDER_MAINTAIN
        - Set default value to 'L1' if not provided
        
        ## 7. Error Handling
        - Validate input values against ZDELIVERY_PRIORITY_TEXT
        - Display error message if invalid value is entered
        
        ## 8. Authorization
        - Standard SD authorization objects apply
        - No additional authorization checks required
        
        ## 9. Testing Strategy
        - Create test sales order with each priority code
        - Verify field appears in correct position
        - Verify value flow to delivery document
        - Test F4 help functionality
        
        ## 10. Transport Strategy
        - Create transport request for development objects
        - Include all objects in single transport
        """
        
        st.text_area("Example Technical Specification", sample_ts, height=200)
        
        if st.button("Use Example"):
            st.session_state.ts_content = sample_ts
            st.rerun()

# Code to Test Page Function
def show_code_to_test():
    """Display the Code to Unit Test automation page."""
    st.header("Code → Unit Test Automation")
    st.markdown("""
        Generate comprehensive unit tests for SAP code. 
        Upload or paste your code below.
    """)
    
    # Initialize session state for this page
    if 'code_content' not in st.session_state:
        st.session_state.code_content = ""
    if 'test_result' not in st.session_state:
        st.session_state.test_result = ""
    if 'show_thinking' not in st.session_state:
        st.session_state.show_thinking = False
    if 'thinking_steps' not in st.session_state:
        st.session_state.thinking_steps = []
    if 'use_agent_framework' not in st.session_state:
        st.session_state.use_agent_framework = False
    if "use_autogen" not in st.session_state:
        st.session_state.use_autogen = False
    
    # File upload or text input
    st.subheader("Input Code")
    
    tab1, tab2 = st.tabs(["Upload File", "Paste Code"])
    
    with tab1:
        uploaded_file = st.file_uploader("Upload a code file", 
                                        type=['txt', 'abap', 'js'],
                                        help="Supported formats: .txt, .abap, .js")
        
        if uploaded_file is not None:
            try:
                # Extract text from the uploaded file
                code_text = DocumentProcessor.extract_text(uploaded_file)
                st.session_state.code_content = code_text
                st.success(f"Successfully loaded code from {uploaded_file.name}")
                
                # Show preview
                with st.expander("Preview code"):
                    st.code(code_text)
            except Exception as e:
                st.error(f"Error processing file: {str(e)}")
    
    with tab2:
        code_text_input = st.text_area("Enter your code", 
                                     height=300,
                                     value=st.session_state.code_content)
        
        if code_text_input != st.session_state.code_content:
            st.session_state.code_content = code_text_input
    
    # Options
    st.subheader("Options")
    
    model_options = ["llama-3.3-70b-versatile", "llama3-8b-8192", "mixtral-8x7b-32768","deepseek-r1-distill-llama-70b"]
    selected_model = st.selectbox("Select LLM Model", model_options)
    
    col1, col2 = st.columns(2)
    
    with col1:
        code_type = st.radio("Code Type", ["ABAP", "BTP/Node.js"], 
                           help="Select the type of code you're testing")
    
    with col2:
        st.session_state.use_agent_framework = st.checkbox("Use Multi-Agent Framework", value=False,
                                                          help="Use multiple specialized agents to generate the tests")
        st.session_state.use_autogen = st.checkbox("Use Microsoft AutoGen Framework", value=False, help="Enables agent-to-agent collaboration using Microsoft AutoGen framework")
    
    # Advanced options
    with st.expander("Advanced Options"):
        st.session_state.show_thinking = st.checkbox("Show agent thinking process", value=False)
        test_coverage = st.slider("Target Test Coverage (%)", 60, 100, 80)
    
    # Generate button
    if st.button("Generate Unit Tests", type="primary", 
                disabled=not st.session_state.code_content):
        
        if not st.session_state.code_content.strip():
            st.error("Please provide code content first.")
            return
        
        with st.spinner("Generating Unit Tests..."):
            try:
                # Initialize GROQ client with selected model
                groq_client = GroqClient()
                groq_client.set_model(selected_model)
                
                # Determine code type for processing
                code_type_value = "ABAP" if code_type == "ABAP" else "Node.js"
                
                if st.session_state.use_autogen:
                    # Use AutoGen integration
                    autogen_integration = AgentFactory.create_autogen_integration(groq_client)
                    
                    # Execute the task
                    result = autogen_integration.execute_task(
                        f"Generate comprehensive unit tests for this {code_type_value} code with {test_coverage}% coverage target.",
                        st.session_state.code_content,
                        "code_to_test"
                    )
                    
                    st.session_state.test_result = result
                    st.session_state.thinking_steps = ["Using Microsoft AutoGen agent collaboration"]
                    
                elif st.session_state.use_agent_framework:
                    # Use multi-agent framework
                    orchestrator = MultiAgentOrchestrator(groq_client)
                    
                    # Execute the task
                    result = orchestrator.execute_task(
                        f"Generate comprehensive unit tests for this {code_type_value} code with {test_coverage}% coverage target.",
                        st.session_state.code_content
                    )
                    
                    st.session_state.test_result = result
                    st.session_state.thinking_steps = orchestrator.get_thinking_steps()
                    
                else:
                    # Use single agent
                    test_generator = AgentFactory.create_test_case_generator(groq_client)
                    st.session_state.test_result = test_generator.generate_unit_tests(
                        st.session_state.code_content,
                        code_type_value
                    )
                    st.session_state.thinking_steps = test_generator.get_thinking_steps()
                
                st.success("Unit tests generated successfully!")
                
            except Exception as e:
                st.error(f"Error generating unit tests: {str(e)}")
    
    # Display results if available
    if st.session_state.test_result:
        st.subheader("Generated Unit Tests")
        
        # Display thinking process if enabled
        if st.session_state.show_thinking and st.session_state.thinking_steps:
            with st.expander("Agent Thinking Process", expanded=False):
                for i, step in enumerate(st.session_state.thinking_steps):
                    st.markdown(f"{i+1}. {step}")
        
        # Display the generated tests
        st.code(st.session_state.test_result)
        
        # Feedback and regeneration
        feedback = st.text_area("Feedback for improvement (optional)")
        
        if feedback and st.button("Improve Tests Based on Feedback"):
            with st.spinner("Improving tests..."):
                try:
                    # Initialize GROQ client with selected model
                    groq_client = GroqClient()
                    groq_client.set_model(selected_model)
                    
                    # Determine code type for processing
                    code_type_value = "ABAP" if code_type == "ABAP" else "Node.js"
                    
                    # Use test generator to improve the tests
                    test_generator = AgentFactory.create_test_case_generator(groq_client)
                    improved_tests = test_generator.improve_tests(
                        st.session_state.test_result,
                        feedback,
                        code_type_value
                    )
                    
                    # Update the result
                    st.session_state.test_result = improved_tests
                    st.session_state.thinking_steps.extend(test_generator.get_thinking_steps())
                    
                    st.success("Tests improved successfully!")
                    st.rerun()
                    
                except Exception as e:
                    st.error(f"Error improving tests: {str(e)}")
        
        # Download options
        st.subheader("Download Options")
        
        test_ext = ".abap" if code_type == "ABAP" else ".test.js"
        
        FileHandler.display_download_button(
            st.session_state.test_result,
            f"unit_tests{test_ext}",
            f"Download Tests ({test_ext})",
            "text/plain"
        )
        
        # Option to create package with code and tests
        with st.expander("Create Complete Package"):
            st.markdown("Create a zip package containing both the original code and unit tests.")
            
            zip_files = {}
            
            if code_type == "ABAP":
                # Add files for ABAP code
                zip_files["original_code.abap"] = st.session_state.code_content
                zip_files["unit_tests.abap"] = st.session_state.test_result
                zip_files["README.md"] = "# ABAP Code with Unit Tests\n\nThis package contains ABAP code and its corresponding unit tests.\n\n## Implementation Notes\n\nPlease ensure your ABAP environment supports ABAP Unit testing."
            else:
                # Add files for BTP/Node.js code
                zip_files["src/index.js"] = st.session_state.code_content
                zip_files["test/index.test.js"] = st.session_state.test_result
                zip_files["README.md"] = "# Node.js Code with Unit Tests\n\nThis package contains Node.js code and its corresponding Jest unit tests.\n\n## Implementation Notes\n\nEnsure you have Jest installed in your environment."
            
            zip_bytes = FileHandler.create_zip_download(zip_files)
            
            FileHandler.display_download_button(
                zip_bytes,
                "code_with_tests.zip",
                "Download Complete Package (.zip)",
                "application/zip"
            )

# Enhancement Ideas Page Function
def show_enhancement_ideas():
    """Display the Enhancement Ideas page."""
    st.header("Enhancement Ideas")
    st.markdown("""
        Get suggestions for improving your existing SAP code. 
        Upload or paste your code below to receive enhancement recommendations.
    """)
    
    # Initialize session state for this page
    if 'code_content' not in st.session_state:
        st.session_state.code_content = ""
    if 'enhancement_result' not in st.session_state:
        st.session_state.enhancement_result = ""
    if 'show_thinking' not in st.session_state:
        st.session_state.show_thinking = False
    if 'thinking_steps' not in st.session_state:
        st.session_state.thinking_steps = []
    if 'use_agent_framework' not in st.session_state:
        st.session_state.use_agent_framework = False
    
    # File upload or text input
    st.subheader("Input Code")
    
    tab1, tab2 = st.tabs(["Upload File", "Paste Code"])
    
    with tab1:
        uploaded_file = st.file_uploader("Upload a code file", 
                                        type=['txt', 'abap', 'js'],
                                        help="Supported formats: .txt, .abap, .js")
        
        if uploaded_file is not None:
            try:
                # Extract text from the uploaded file
                code_text = DocumentProcessor.extract_text(uploaded_file)
                st.session_state.code_content = code_text
                st.success(f"Successfully loaded code from {uploaded_file.name}")
                
                # Show preview
                with st.expander("Preview code"):
                    st.code(code_text)
            except Exception as e:
                st.error(f"Error processing file: {str(e)}")
    
    with tab2:
        code_text_input = st.text_area("Enter your code", 
                                     height=300,
                                     value=st.session_state.code_content)
        
        if code_text_input != st.session_state.code_content:
            st.session_state.code_content = code_text_input
    
    # Options
    st.subheader("Options")
    
    model_options = ["llama-3.3-70b-versatile", "llama3-8b-8192", "mixtral-8x7b-32768","deepseek-r1-distill-llama-70b"]
    selected_model = st.selectbox("Select LLM Model", model_options)
    
    col1, col2 = st.columns(2)
    
    with col1:
        code_type = st.radio("Code Type", ["ABAP", "BTP/Node.js"], 
                           help="Select the type of code you're enhancing")
    
    if "use_autogen" not in st.session_state:
        st.session_state.use_autogen = False
    with col2:
        enhancement_focus = st.multiselect(
            "Enhancement Focus Areas",
            ["Performance", "Security", "Maintainability", "Modernization", "Cloud Readiness"],
            ["Performance", "Maintainability"]
        )
    
    # Advanced options
    with st.expander("Advanced Options"):
        st.session_state.show_thinking = st.checkbox("Show agent thinking process", value=False)
        st.session_state.use_autogen = st.checkbox("Use Microsoft AutoGen Framework", value=False, help="Enables agent-to-agent collaboration using Microsoft AutoGen framework")
        st.session_state.use_agent_framework = st.checkbox("Use Multi-Agent Framework", value=False,
                                                         help="Use multiple specialized agents to suggest enhancements")
    
    # Generate button
    if st.button("Suggest Enhancements", type="primary", 
                disabled=not st.session_state.code_content):
        
        if not st.session_state.code_content.strip():
            st.error("Please provide code content first.")
            return
            if st.session_state.use_autogen:
                    # Use AutoGen integration
                    autogen_integration = AgentFactory.create_autogen_integration(groq_client)
                    
                    # Execute the task
                    result = autogen_integration.execute_task(
                        f"Suggest enhancements for this {code_type_value} code. {focus_areas_text}",
                        st.session_state.code_content,
                        "enhancement_ideas"
                    )
                    
                    st.session_state.enhancement_result = result
                    st.session_state.thinking_steps = ["Using Microsoft AutoGen agent collaboration"]

        with st.spinner("Analyzing code and suggesting enhancements..."):
            try:
                # Initialize GROQ client with selected model
                groq_client = GroqClient()
                groq_client.set_model(selected_model) 
                # Determine code type for processing
                code_type_value = "ABAP" if code_type == "ABAP" else "BTP"
                # Add focus areas to the prompt if selected
                focus_areas_text = ""
                if enhancement_focus:
                    focus_areas_text = f"Focus primarily on: {', '.join(enhancement_focus)}"
                if st.session_state.use_agent_framework:
                    # Use multi-agent framework
                    orchestrator = MultiAgentOrchestrator(groq_client)
                    # Execute the task
                    result = orchestrator.execute_task(
                        f"Suggest enhancements for this {code_type_value} code. {focus_areas_text}",
                        st.session_state.code_content
                    )
                    st.session_state.enhancement_result = result
                    st.session_state.thinking_steps = orchestrator.get_thinking_steps()
                else:
                    # Use single agent
                    enhancement_suggester = AgentFactory.create_enhancement_suggester_agent(groq_client)
                    st.session_state.enhancement_result = enhancement_suggester.suggest_enhancements(
                        st.session_state.code_content,
                        code_type_value
                    )
                    st.session_state.thinking_steps = enhancement_suggester.get_thinking_steps()
                st.success("Enhancement suggestions generated successfully!")
            except Exception as e:
                st.error(f"Error generating enhancement suggestions: {str(e)}")
    # Display results if available
    if st.session_state.enhancement_result:
        st.subheader("Enhancement Suggestions")
        # Display thinking process if enabled
        if st.session_state.show_thinking and st.session_state.thinking_steps:
            with st.expander("Agent Thinking Process", expanded=False):
                for i, step in enumerate(st.session_state.thinking_steps):
                    st.markdown(f"{i+1}. {step}")
        # Display the enhancement suggestions
        st.markdown(st.session_state.enhancement_result)
        # Download options
        st.subheader("Download Options")
        FileHandler.display_download_button(
            st.session_state.enhancement_result,
            "enhancement_suggestions.md",
            "Download Suggestions (.md)",
            "text/markdown"
        )
        # Create Word document download
        docx_bytes = FileHandler.create_docx_download(st.session_state.enhancement_result)
        FileHandler.display_download_button(
            docx_bytes,
            "enhancement_suggestions.docx",
            "Download as Word (.docx)",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# Transport Doc Page Function
def show_transport_docs():
    """Display the Transport Request Documentation page."""
    st.header("Transport Request Documentation")
    st.markdown("""
        Generate comprehensive documentation for SAP transport requests.
        Upload or paste your code, TS, or FS to create standardized transport documentation.
    """)
    # Initialize session state for this page
    if 'content' not in st.session_state:
        st.session_state.content = ""
    if 'documentation_result' not in st.session_state:
        st.session_state.documentation_result = ""
    if 'show_thinking' not in st.session_state:
        st.session_state.show_thinking = False
    if 'thinking_steps' not in st.session_state:
        st.session_state.thinking_steps = []
    if 'transport_id' not in st.session_state:
        st.session_state.transport_id = ""
    if 'use_agent_framework' not in st.session_state:
        st.session_state.use_agent_framework = False
    # File upload or text input
    st.subheader("Input Content")
    tab1, tab2 = st.tabs(["Upload File", "Paste Content"])
    with tab1:
        uploaded_file = st.file_uploader("Upload a file", 
                                        type=['txt', 'pdf', 'docx', 'abap', 'js'],
                                        help="Supported formats: .txt, .pdf, .docx, .abap, .js")
        if uploaded_file is not None:
            try:
                # Extract text from the uploaded file
                content_text = DocumentProcessor.extract_text(uploaded_file)
                st.session_state.content = content_text
                st.success(f"Successfully loaded content from {uploaded_file.name}")
                # Show preview
                with st.expander("Preview content"):
                    st.text_area("File content", content_text, height=200)
            except Exception as e:
                st.error(f"Error processing file: {str(e)}")
    with tab2:
        content_text_input = st.text_area("Enter your content", 
                                       height=300,
                                       value=st.session_state.content)
        if content_text_input != st.session_state.content:
            st.session_state.content = content_text_input
    # Options
    st.subheader("Options")
    model_options = ["llama-3.3-70b-versatile", "llama3-8b-8192", "mixtral-8x7b-32768","deepseek-r1-distill-llama-70b"]
    selected_model = st.selectbox("Select LLM Model", model_options)
    transport_id=""
    col1, col2 = st.columns(2)
    with col1:
        content_type = st.radio("Content Type", ["Code", "Technical Specification", "Functional Specification"], 
                              help="Select the type of content you're providing")
        content_type_desc = {
            "Code": "code change",
            "Technical Specification": "technical specification",
            "Functional Specification": "functional specification"
        }[content_type]
    with col2:
        if "use_autogen" not in st.session_state:
            st.session_state.use_autogen = False
            transport_id = st.text_input("Transport Request ID (optional)", 
                                  value=st.session_state.transport_id,
                                  help="Enter the SAP transport request ID if known")
        if transport_id != st.session_state.transport_id:
            st.session_state.transport_id = transport_id
    # Transport type
    transport_type = st.radio("Transport Type", ["Workbench", "Customizing"], horizontal=True)
    # System path
    system_path = st.selectbox("System Path",
                            ["DEV → QA → PROD",
                             "DEV → TEST → QA → PROD",
                             "SANDBOX → DEV → PROD",
                             "DEV → UAT → PROD"])
    st.session_state.use_autogen = st.checkbox("Use Microsoft AutoGen Framework", value=False, help="Enables agent-to-agent collaboration using Microsoft AutoGen framework")
    # Advanced options
    with st.expander("Advanced Options"):
        st.session_state.show_thinking = st.checkbox("Show agent thinking process", value=False)
        st.session_state.use_agent_framework = st.checkbox("Use Multi-Agent Framework", value=False,
                                                         help="Use multiple specialized agents to generate documentation")
    # Generate button
    if st.session_state.use_autogen:
                    # Use AutoGen integration
                    autogen_integration = AgentFactory.create_autogen_integration(groq_client)
                    # Execute the task
                    result = autogen_integration.execute_task(
                        f"Generate a transport request document for this SAP {content_type_desc}.",
                        st.session_state.content,
                        "transport_docs"
                    )
                    st.session_state.doc_result = result
                    st.session_state.thinking_steps = ["Using Microsoft AutoGen agent collaboration"]

    if st.button("Generate Transport Documentation", type="primary", 
                disabled=not st.session_state.content):
        if not st.session_state.content.strip():
            st.error("Please provide content first.")
            return
        with st.spinner("Generating transport documentation..."):
            try:
                # Initialize GROQ client with selected model
                groq_client = GroqClient()
                groq_client.set_model(selected_model)
                # Determine content type for processing
                if content_type == "Code":
                    content_type_value = "CODE"
                elif content_type == "Technical Specification":
                    content_type_value = "TS"
                else:
                    content_type_value = "FS"
                # Add transport details to the prompt
                transport_details = f"""
                Transport Request ID: {st.session_state.transport_id if st.session_state.transport_id else '[To be filled by user]'}
                Transport Type: {transport_type}
                System Path: {system_path}
                """
                enhanced_content = f"{transport_details}\n\n{st.session_state.content}"
                if st.session_state.use_agent_framework:
                    # Use multi-agent framework
                    orchestrator = MultiAgentOrchestrator(groq_client)
                    # Execute the task
                    result = orchestrator.execute_task(
                        f"Generate transport request documentation for this {content_type_value} with the provided transport details.",
                        enhanced_content
                    )
                    st.session_state.documentation_result = result
                    st.session_state.thinking_steps = orchestrator.get_thinking_steps()
                else:
                    # Use single agent
                    transport_doc_generator = AgentFactory.create_transport_doc_generator(groq_client)
                    st.session_state.documentation_result = transport_doc_generator.generate_transport_doc(
                        enhanced_content,
                        content_type_value
                    )
                    st.session_state.thinking_steps = transport_doc_generator.get_thinking_steps()
                st.success("Transport documentation generated successfully!")
            except Exception as e:
                st.error(f"Error generating transport documentation: {str(e)}")
    # Display results if available
    if st.session_state.documentation_result:
        st.subheader("Transport Request Documentation")
        # Display thinking process if enabled
        if st.session_state.show_thinking and st.session_state.thinking_steps:
            with st.expander("Agent Thinking Process", expanded=False):
                for i, step in enumerate(st.session_state.thinking_steps):
                    st.markdown(f"{i+1}. {step}")
        # Display the generated documentation
        st.markdown(st.session_state.documentation_result)
        # Download options
        st.subheader("Download Options")
        # Text download
        FileHandler.display_download_button(
            st.session_state.documentation_result,
            "transport_documentation.md",
            "Download as Markdown (.md)",
            "text/markdown"
        )
        # Word document download
        docx_bytes = FileHandler.create_docx_download(st.session_state.documentation_result)
        FileHandler.display_download_button(
            docx_bytes,
            "transport_documentation.docx",
            "Download as Word (.docx)",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

###################
# Main Application
###################

# Initialize session state for API key
if 'groq_api_key' not in st.session_state:
    st.session_state.groq_api_key = os.getenv("GROQ_API_KEY", "")
if 'api_key_saved' not in st.session_state:
    st.session_state.api_key_saved = False

# Page configuration
st.set_page_config(
    page_title="SAP AI Assistant – Automating Functional to Technical to Code",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Title and description
st.title("SAP AI Assistant – Automating Functional to Technical to Code")
st.markdown("""
    <div style='background-color: #f0f2f6; padding: 1em; border-radius: 10px;'>
    <p>This AI-powered assistant helps SAP consultants and developers automate workflows from 
    functional specifications to technical design, code generation, and testing.</p>
    </div>
""", unsafe_allow_html=True)

# Sidebar navigation
with st.sidebar:
    st.image("https://www.sap.com/dam/application/shared/logos/sap-logo-svg.svg", width=150)
    st.markdown("## SAP AI Assistant")
    selected = option_menu(
        menu_title=None,
        options=[
            "FS → TS Automation", 
            "TS → Code Automation", 
            "Code → Unit Test", 
            "Enhancement Ideas", 
            "Transport Request Docs"
        ],
        icons=[
            "file-earmark-text", 
            "code-square", 
            "check-circle", 
            "lightbulb", 
            "truck"
        ],
        menu_icon="cast",
        default_index=0,
        styles={
            "container": {"padding": "0!important", "background-color": "#f0f2f6"},
            "icon": {"color": "#0F5BA7", "font-size": "14px"},
            "nav-link": {"font-size": "14px", "text-align": "left", "margin": "0px", "--hover-color": "#e6e6e6"},
            "nav-link-selected": {"background-color": "#0F5BA7"}
        }
    )
    # GROQ API Key Configuration
    st.markdown("### API Configuration")
    with st.expander("Groq API Settings", expanded=not st.session_state.api_key_saved):
        api_key = st.text_input(
            "Groq API Key", 
            value=st.session_state.groq_api_key,
            type="password",
            help="Enter your Groq API key to use the LLM services"
        )
        if st.button("Save API Key"):
            if api_key:
                st.session_state.groq_api_key = api_key
                os.environ["GROQ_API_KEY"] = api_key
                st.session_state.api_key_saved = True
                st.success("API key saved successfully!")
            else:
                st.error("Please enter a valid API key")
    st.markdown("### About")
    st.info(
        """
        This application uses Generative AI and Multi-Agent collaboration to automate SAP workflows.
        """
    )
    st.markdown("### Help")
    with st.expander("How to use this app"):
        st.markdown("""
            - Enter your Groq API key in the sidebar
            - Select a task from the menu
            - Upload or paste your input text
            - Configure options
            - Click generate and wait for results
            - Download the output
        """)
    # App footer
    st.markdown("---")
    st.markdown("© 2023 SAP AI Assistant")

# Route to the selected page
if selected == "FS → TS Automation":
    show_fs_to_ts()
elif selected == "TS → Code Automation":
    show_ts_to_code()
elif selected == "Code → Unit Test":
    show_code_to_test()
elif selected == "Enhancement Ideas":
    show_enhancement_ideas()
elif selected == "Transport Request Docs":
    show_transport_docs()